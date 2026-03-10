"""
ingestor.py — GOIES Ingestion Engine
Handles URL scraping, PDF parsing, DOCX parsing.
All functions return plain UTF-8 text ready for extraction.
"""

from __future__ import annotations

import re
from typing import Optional

MAX_TEXT_CHARS = 50_000  # cap to avoid overwhelming the LLM


# ── URL Ingestion ─────────────────────────────────────────────────────────────


def fetch_url_text(url: str, timeout: int = 15) -> str:
    """
    Fetch a URL and extract readable article text.
    Requires: httpx, readability-lxml (preferred) or beautifulsoup4 (fallback).
    Raises: ValueError on bad URL, RuntimeError on fetch/parse failure.
    """
    if not url.startswith(("http://", "https://")):
        raise ValueError(f"Invalid URL (must start with http:// or https://): {url}")

    try:
        import httpx
    except ImportError:
        raise RuntimeError("httpx not installed. Run: pip install httpx")

    try:
        with httpx.Client(
            follow_redirects=True,
            timeout=timeout,
            headers={
                "User-Agent": (
                    "Mozilla/5.0 (compatible; GOIES/3.0; +https://github.com/tanu-1403/GOIES)"
                )
            },
        ) as client:
            resp = client.get(url)
            resp.raise_for_status()
            html = resp.text
    except httpx.HTTPStatusError as e:
        raise RuntimeError(f"HTTP {e.response.status_code} fetching URL: {url}")
    except httpx.RequestError as e:
        raise RuntimeError(f"Network error fetching URL: {e}")

    text = _extract_text_from_html(html, url)
    return text[:MAX_TEXT_CHARS]


def _extract_text_from_html(html: str, url: str = "") -> str:
    """Try readability-lxml first, fall back to BeautifulSoup."""
    # Attempt 1: readability-lxml (best quality)
    try:
        from readability import Document

        doc = Document(html)
        content_html = doc.summary()
        # Strip HTML tags from readability output
        clean = re.sub(r"<[^>]+>", " ", content_html)
        clean = re.sub(r"\s{2,}", " ", clean).strip()
        if len(clean) > 200:
            return clean
    except ImportError:
        pass
    except Exception:
        pass

    # Attempt 2: BeautifulSoup
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, "html.parser")
        # Remove boilerplate tags
        for tag in soup(["script", "style", "nav", "header", "footer", "aside", "ads"]):
            tag.decompose()
        # Try article/main first, fall back to body
        container = (
            soup.find("article")
            or soup.find("main")
            or soup.find(id=re.compile(r"(content|article|main)", re.I))
            or soup.find("body")
        )
        if container:
            text = container.get_text(separator=" ", strip=True)
        else:
            text = soup.get_text(separator=" ", strip=True)
        clean = re.sub(r"\s{2,}", " ", text).strip()
        if len(clean) > 100:
            return clean
    except ImportError:
        pass
    except Exception:
        pass

    # Attempt 3: Crude tag strip
    clean = re.sub(r"<[^>]+>", " ", html)
    clean = re.sub(r"\s{2,}", " ", clean).strip()
    if len(clean) < 100:
        raise RuntimeError(
            "Could not extract readable text from URL. Install readability-lxml or beautifulsoup4."
        )
    return clean


# ── PDF Parsing ───────────────────────────────────────────────────────────────


def parse_pdf(content: bytes) -> str:
    """
    Extract text from a PDF byte payload.
    Tries pypdf first (maintained fork), falls back to PyPDF2.
    Raises RuntimeError if neither is installed.
    """
    # Attempt 1: pypdf (modern maintained fork)
    try:
        import pypdf

        reader = pypdf.PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                continue
        text = "\n".join(pages).strip()
        if text:
            return text[:MAX_TEXT_CHARS]
    except ImportError:
        pass

    # Attempt 2: PyPDF2 (legacy)
    try:
        import PyPDF2
        import io

        reader = PyPDF2.PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                continue
        text = "\n".join(pages).strip()
        if text:
            return text[:MAX_TEXT_CHARS]
    except ImportError:
        pass

    raise RuntimeError("PDF parsing requires pypdf or PyPDF2. Run: pip install pypdf")


# ── DOCX Parsing ──────────────────────────────────────────────────────────────


def parse_docx(content: bytes) -> str:
    """
    Extract text from a DOCX byte payload.
    Requires python-docx.
    Raises RuntimeError if not installed.
    """
    try:
        import io
        import docx

        doc = docx.Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also capture tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n".join(paragraphs).strip()
        return text[:MAX_TEXT_CHARS]
    except ImportError:
        raise RuntimeError(
            "DOCX parsing requires python-docx. Run: pip install python-docx"
        )
    except Exception as e:
        raise RuntimeError(f"Failed to parse DOCX: {e}")


import io
